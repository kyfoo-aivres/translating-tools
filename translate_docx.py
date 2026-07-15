import argparse
import os
import shutil
import glob
from docx import Document
from googletrans import Translator
from tqdm import tqdm
import asyncio

def extract_runs(document):
    """
    Extract all runs from the document, including from paragraphs, tables, and shapes (text boxes).
    """
    all_runs = []
    
    # Extract from paragraphs
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.text and run.text.strip():
                all_runs.append(run)
    
    # Extract from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text and run.text.strip():
                            all_runs.append(run)
    
    # Extract from shapes (text boxes, etc.) if available
    try:
        if hasattr(document, 'shapes'):
            for shape in document.shapes:
                if hasattr(shape, 'text_frame') and shape.text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        for run in paragraph.runs:
                            if run.text and run.text.strip():
                                all_runs.append(run)
    except Exception as e:
        print(f"Note: Could not extract text from shapes: {e}")
    
    return all_runs

def translate_document(input_file, src_lang="zh-cn", dest_lang="en"):
    # Ensure input file exists
    if not os.path.exists(input_file):
        print(f"Error: The file '{input_file}' was not found.")
        return

    # Prepare input, output, and log paths
    file_directory = os.path.dirname(os.path.abspath(input_file))
    base_name = os.path.splitext(os.path.basename(input_file))[0]
    ext_name = os.path.splitext(input_file)[1]
    
    # Setup paths according to the requirements
    output_file = os.path.join(file_directory, f"{base_name}_{dest_lang}{ext_name}")
    issue_file = os.path.join(file_directory, f"Issues_{base_name}.txt")

    # Duplicate the source DOCX
    try:
        shutil.copy2(input_file, output_file)
        print(f"Copied source file to: {output_file}")
    except Exception as e:
        print(f"Error copying the file: {e}")
        return

    # Load the duplicated DOCX
    try:
        doc = Document(output_file)
    except Exception as e:
        print(f"Error opening copied document (make sure it's a valid .docx): {e}")
        return

    # Extract all textual components that need to be translated
    print("Analyzing document and extracting texts...")
    all_runs = extract_runs(doc)
    
    print(f"Discovered {len(all_runs)} text runs to translate.")
    
    # Setup issue logging 
    if os.path.exists(issue_file):
        os.remove(issue_file) # Clean up old issues file if it exists
        
    issue_found = False

    async def translate_run_list(runs_list, desc):
        nonlocal issue_found
        # Initialize the translator inside the event loop
        translator = Translator()
        
        # Proceed with the translation and show the progress bar using tqdm
        for run in tqdm(runs_list, desc=desc, unit="run"):
            original_text = run.text
            try:
                # Perform the translation
                translation = translator.translate(original_text, src=src_lang, dest=dest_lang)
                
                # Wait for the result if it's an async coroutine
                if asyncio.iscoroutine(translation):
                    translation = await translation
                    
                # Update the run's text with the translated text
                if translation and hasattr(translation, 'text'):
                    run.text = translation.text
            except Exception as e:
                issue_found = True
                # Log the issue into the issue file
                with open(issue_file, "a", encoding="utf-8") as file:
                    file.write(f"Failed to translate block ({desc}): '{original_text}'\n")
                    file.write(f"Error details: {str(e)}\n")
                    file.write("-" * 50 + "\n")
                    
    async def main_translation_task():
        if all_runs:
            await translate_run_list(all_runs, "Translating Document")

    # Execute the async translation process
    asyncio.run(main_translation_task())
    
    # Post-translation: Fix missing spaces between runs in Latin-based translations
    cjk_langs = ['zh', 'zh-cn', 'zh-tw', 'ja', 'ko']
    if dest_lang.lower() not in cjk_langs:
        no_space_before = tuple('.,!?:;)]}”’')
        no_space_after = tuple('([{"‘')
        
        # Collect all paragraphs from document, tables, and shapes
        all_paragraphs = []
        all_paragraphs.extend(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)
        try:
            if hasattr(doc, 'shapes'):
                for shape in doc.shapes:
                    if hasattr(shape, 'text_frame') and shape.text_frame:
                        all_paragraphs.extend(shape.text_frame.paragraphs)
        except Exception as e:
            pass  # Silently skip if shapes cannot be accessed
        
        for paragraph in all_paragraphs:
            runs = paragraph.runs
            if not runs:
                continue
                
            for i in range(len(runs) - 1):
                curr_run = runs[i]
                next_run = runs[i+1]
                
                c_text = curr_run.text
                n_text = next_run.text
                
                if not c_text or not n_text:
                    continue
                    
                # If current run doesn't end with a space, and next doesn't start with a space
                if not c_text[-1].isspace() and not n_text[0].isspace():
                    # Add a space to current_run unless it breaks punctuation spacing rules
                    if not n_text.startswith(no_space_before) and not c_text.endswith(no_space_after):
                        curr_run.text = c_text + " "
                
    # Save the updated document
    try:
        doc.save(output_file)
        print(f"\nTranslation complete. The translated file is ready at: {output_file}")
        print("Note: If the document contains a table of contents, please update it in Microsoft Word (right-click TOC > Update Field > Update entire table).")
    except Exception as e:
        print(f"\nError: Could not save the translated document: {e}")

    # Notify the user if issues were logged
    if issue_found:
        print(f"Note: Some issues occurred during translation. See '{issue_file}' for details.")

def process_file(file_path, src_lang, dest_lang):
    print(f"\n{'='*50}\nProcessing file: {file_path}\n{'='*50}")
    translate_document(file_path, src_lang, dest_lang)

def main():
    parser = argparse.ArgumentParser(description="Translate Microsoft Word (.docx) files.")
    
    # Mutually exclusive group for file or directory input
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("-f", "--file", help="Path to a single source DOCX file to translate")
    group.add_argument("-d", "--directory", help="Path to a directory containing DOCX files to translate")
    
    parser.add_argument("--source", default="zh-cn", help="Source language code (default: zh-CN)")
    parser.add_argument("--target", default="en", help="Target language code (default: en)")
    
    args = parser.parse_args()
    
    if args.file:
        process_file(args.file, args.source, args.target)
    elif args.directory:
        if not os.path.isdir(args.directory):
            print(f"Error: Directory '{args.directory}' does not exist.")
            return
            
        docx_files = glob.glob(os.path.join(args.directory, "*.docx"))
        
        if not docx_files:
            print(f"No .docx files found in directory: {args.directory}")
            return
            
        print(f"Discovered {len(docx_files)} .docx files in '{args.directory}'.")
        for docx in docx_files:
            # Skip hidden/temporary Word files (though less common, just in case)
            if not os.path.basename(docx).startswith("~$"):
                process_file(docx, args.source, args.target)

if __name__ == "__main__":
    main()